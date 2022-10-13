# -*- coding: utf-8 -*-
"""
Created on 13/10/2022

@author: Davide laghi

A file named 'Acronyms Table.docx" will be generated that will contain the 
acronym table.
"""
# ################ IMPORT ####################################################
import docx  # to be downloaded from conda-forge
import re
import os
import logging

from typing import Tuple

# ################ USER INPUT ################################################
# Path to the target folder containing target files
FOLDER = r'C:\Users\d.laghi\Desktop\chapters'
# path to the output folder
OUTPATH = r'C:\Users\d.laghi\Desktop\test'
read_tables = False  # Boolean (True/False) to scan or not tables
try_definitions = True  # If True tries to find the acronym definition
MODE = 'latex'  # Check ALLOWED_MODES

# ############### EXPERT PARAMETERS ##########################################
# --- Acronym patterns ---
# 2+ Uppercase letters (e.g. DT, CHT)
PAT_ACRONYM = re.compile(r'[A-Z][A-Z][A-Z]*')
# Blank line
PAT_BLANK = re.compile(r'^\s*$')
# # Match an acronym only if between brackets (e.g. (DT), (CHT))
# PAT_ACRONYM = re.compile(r'(?<=\()[A-Z][A-Z][A-Z]*')
ALLOWED_MODES = ['latex', 'word']  # execution modes allowed

# ################# FUNCTIONS ################################################
def get_acronym(text):
    """
    Given a text, it returns a list of acronyms. An acronym is defined as a
    string composed by all upper case letters (at least of lenght two).

    Parameters
    ----------
    text : str
        text to be analyzed.

    Returns
    -------
    acronyms: list
        list of acronyms

    """
    acronyms = PAT_ACRONYM.findall(text)

    return acronyms


def get_definition(text, acronym):
    """
    Check the text for the acronym definition

    Parameters
    ----------
    text : str
        text were to search the definition.
    acronym : str
        acronym to search.

    Returns
    -------
    None.

    """
    onechar = '\w+[\s-]+'
    pat_string = ''
    for char in acronym:
        pat_string = pat_string+char+onechar

    pat = re.compile(pat_string)

    try:
        match = pat.search(text).group()
    except AttributeError:
        match = None

    return match

def get_acronyms_from_file(file: os.PathLike, mode: str) -> Tuple[list]:
    """returns acronyms found in a file

    Parameters
    ----------
    file : os.PathLike
        path to the target file
    mode : str
        execution mode

    Returns
    -------
    acronyms_sorted: list
        acronyms found in file
    
    paragraphs: list
        list of paragraphs in the document

    Raises
    ------
    ValueError
        if execution mode is not available
    """
    # --- Read the acronyms ---
    acronyms = []  # list of acronyms

    # Open the document and collect paragraphs
    if mode == 'word':
        document = docx.Document(file)
        print('Scanning the text for acronyms...')
        # iterate on document paragraph
        paragraphs = []
        for par in document.paragraphs:
            # iterate on paragraph characters
            paragraphs.append(par.text)

    elif mode == 'latex':
        with open(file, 'r') as infile:
            # a paragraph is delimited in latex by a blank line
            paragraphs = []
            par = ''
            try:
                for line in infile:
                    # Blank line is found, old paragraph to be saved and new
                    # one to be created
                    if PAT_BLANK.match(line) is not None:
                        paragraphs.append(par)
                        par = ''
                    else:
                    # if no blank line just keep on building the paragraph
                        par = par+line
            except UnicodeDecodeError as e:
                logging.warning('Line skipped: {}'.format(e))
    else:
        raise ValueError('Selected executions mode is not available. Allowed modes are {}'.format(ALLOWED_MODES))
            
    # Cycle on all paragraphs
    for par in paragraphs:
        # iterate on paragraph characters
        new_acronyms = get_acronym(par)  # text scanner
        acronyms.extend(new_acronyms)

    if read_tables and mode == 'word':  # has to be activated from user inputs
        print('Scanning the tables for acronyms...')
        # iterazione sulle tabelle
        for table in document.tables:
            # Iterazione sulle righe delle tabelle
            for row in table.rows:
                # Iterazione sulle celle
                for cell in row.cells:
                    new_acronyms = get_acronym(cell.text)
                    acronyms.extend(new_acronyms)
        
    # Reoder and avoid duplicates
    acronyms_sorted = sorted(list(set(acronyms)))
    
    return acronyms_sorted, paragraphs


def check_definitions(acronyms: list, paragraphs: list) -> dict:
    # Check for definitions
    definitions = {}
    print('Checking for definitions...')
    for acronym in acronyms:
        pat = re.compile(acronym)
        added = False  # always reset
        # Check if definition is available
        for par in paragraphs:
            if pat.search(par) is not None:
                definition = get_definition(par, acronym)
                if definition is not None:
                    definitions[acronym] = definition
                    added = True
                    break
        # If no definition was found
        if added is False:
            definitions[acronym] = None
    
    return definitions

# ###################### CODE ################################################
# --- Compile acronyms table ---
acronyms = []
definitions = {}
for file in os.listdir(FOLDER):
    filepath = os.path.join(FOLDER, file)
    found_acronyms, paragraphs = get_acronyms_from_file(filepath, MODE)
    # get the new acronyms
    new_acronyms = []
    for acronym in found_acronyms:
        if acronym not in acronyms:
            acronyms.append(acronym)
            new_acronyms.append(acronym)

    # check definitions only for the new acronyms
    new_definitions = check_definitions(new_acronyms, paragraphs)
    definitions.update(new_definitions)

# Reoder and avoid duplicates
acronyms_sorted = sorted(list(set(acronyms)))

# Generate the word file with the acronym table
table_acronyms = docx.Document()
table = table_acronyms.add_table(rows=1, cols=2)
header = table.rows[0].cells
header[0].text = 'Acronym'
header[1].text = 'Description'

# Fill the table
print('Filling table...')
for acronym in acronyms_sorted:
    row_cells = table.add_row().cells
    row_cells[0].text = acronym
    definition = definitions[acronym]

    if definition is not None:
        row_cells[1].text = definition

# Save file
outpath = os.path.join(OUTPATH, 'Acronyms Table.docx')
table_acronyms.save(outpath)

print('All done!')
